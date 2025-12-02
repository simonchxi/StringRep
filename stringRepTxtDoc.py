# -*- coding: utf-8 -*-
"""
@time: 2021-11-08
@author: simon


""" 
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

#########定义读文件内容 -->used by replace function
def read_file(infile):
    with open(infile, encoding='UTF-8') as f:
        read_all = f.read()
        f.close()
        print ("readfile done!")  

    return read_all

#########定义写内容到文件 --> used by replace function
def rewrite_file(outfile, data):
    with open(outfile, 'w', encoding='UTF-8') as f:
        f.write(data)
        f.close()
        print ("rewrite file done!")  


########### replace function 传入文件(file),将旧内容(old_content)替换为新内容(new_content)
def replace(infile, outfile):
	# Read the original file
    content = read_file(infile)
    # Split the content into paragraphs
    paragraphs = content.split('\n\
\n\
')
    
    # Dictionary for direct translations
    translateDict = {
            'Market Type Indicator ': '',
            'VANCOUVER DOWNTOWN ': '温哥华市中心',
            'DOWNTOWN ': '市中心', 
            'VANCOUVER WESTSIDE ': '温西',
            'WESTSIDE ': '温西',
            'VANCOUVER EASTSIDE ': '温东',
            'EASTSIDE ': '温东',
            'NORTH VANCOUVER ': '北温',
            'N.VANCOUVER ': '北温',
            'WEST VANCOUVER ': '西温哥华',
            'RICHMOND ': '列治文',
            'TSAWWASSEN ': '特瓦森',
            'LADNER ': '拉得纳',
            'S SURREY WHITE ROCK ': '南萨里白石镇',
            'ATTACHED:': '公寓房屋交易市场类型为：',
            'DETACHED:': '独栋房屋交易市场类型为：',
            'Sellers Market at ': '卖方市场，销售比率为',
            'Balanced Market at ': '均衡市场，销售比率为',
            'Buyers market at ': '买方市场，销售比率为',
            'Buyers Market at ': '买方市场，销售比率为',
            'Sales Ratio average (': '（',
        #    ' homes selling rate)': '套中成交@@@套房）。',
            'Homes are selling on average 100% of list price': '成交价与挂牌价持平。',
        #    'Homes are selling on average ': '成交价比挂牌价',
        #    ' below list price': '@@@低',
        #    ' above list price': '@@@高',
            'Most Active Price Band** ': '最活跃的价格区间：',
            '0,000 to ': '万-',
            ' mil to ': '百万-',
            ' m il to ': '百万-',
            '0,000, ': '万之间，位于',
            ' mil, ': '百万之间，位于',
            ' mil with average ': '百万之间，销售比率',
            ' m il, ': '百万之间，位于',
            ' m il with average ': '百万之间，销售比率',
            '0,000 with average ': '万之间，销售比率',
            ' mil':'百万。',
            ' and up to ': '区域的最大',
            ' and minimum ': '区域的最小',
            ' Sales Ratio (Sellers market)': '（卖方市场）。',
            ' Sales Ratio (Buyers market)': '（买方市场）。',
            ' Sales Ratio (Balenced market)': '（均衡市场）。',
            'Buyers Best Bet** Homes between ': '买家的最佳选择：购买总价在',
            'Buyers Best Bet** Homes in ': '买家的最佳选择：购买位于',
            'Buyers Best Bet** Homes minimum ': '买家的最佳选择：购买最低',
            'Sellers Best Bet** Selling homes in ': '卖家的最佳选择：出售位于',
            ' bedroom properties': '卧室房产。',
            ' bedrooms': '卧室房产。',
            'Buyers Best Bet**': '买家的最佳选择：',
            'Sellers Best Bet** ': '卖家的最佳选择：',
            'Insufficient data': '数据不足。',
            ' and ': '区域的',
            ' to ': '到',
            '： ': '：'
            }

    # Prepare the output content
    output_content = ''

    for paragraph in paragraphs:
        # Translate the paragraph
        translated_paragraph = paragraph
        for EngTxt, ChnTxt in translateDict.items():
            translated_paragraph = translated_paragraph.replace(EngTxt, ChnTxt)
            translated_paragraph = re.sub(r'Homes are selling on average (\d+)% above list price', r'成交价比挂牌价高\1%。', translated_paragraph)
            translated_paragraph = re.sub(r'Homes are selling on average (\d+)% below list price', r'成交价比挂牌价低\1%。', translated_paragraph)
            translated_paragraph = re.sub(r'(\d+(\.\d)?) in (\d+) homes selling rate', r'\3套中成交\1套房', translated_paragraph)
            translated_paragraph = re.sub(r'but with (\d+) sales price band of', r' \1宗交易成交价区间在', translated_paragraph)

        
        # Combine the original and translated paragraphs
        output_content += paragraph + '\n\
' + translated_paragraph + '\n\
\n\
'
    
    # Write the result to the output file
    rewrite_file(outfile, output_content)

    print("replace content in txt done!")


########### define txt file to word function #######
def txt_to_word(input_txt, output_docx):
    # Create a new Word document
    document = Document()
    
    # Open and read the text file
    with open(input_txt, 'r', encoding='utf-8') as file:
        lines = file.readlines()
        for line in lines:
            # Add each line as a paragraph to the Word document
            document.add_paragraph(line.strip())

    for paragraph in document.paragraphs:
        # Access the format of the paragraph
        paragraph_format = paragraph.paragraph_format
        
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Set space after paragraphs to 0 points
        paragraph_format.space_after = Pt(0)

    # Save the document
    document.save(output_docx)
    print(f"Content from {input_txt} has been copied to {output_docx}.")

file_name = '2025 11 November Chinese'
# call function to replace/translate
replace(f'{file_name}.txt', f'{file_name} Translated.txt')

# Use the function to copy contents into docx
txt_to_word(f'{file_name} Translated.txt', f'{file_name} Translated.docx')